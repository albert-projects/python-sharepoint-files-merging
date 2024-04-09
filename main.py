import sys
import tkinter as tk
import tkinter.ttk as ttk
from tkinter import filedialog
from tkinter import messagebox
import os
import re
import threading
from PIL import Image, ImageTk, ImageSequence
from office365.runtime.auth.authentication_context import AuthenticationContext
from office365.sharepoint.client_context import ClientContext
from office365.runtime.client_request_exception import ClientRequestException
from docxcompose.composer import Composer
from docx import Document as Document_compose
from docx import Document
from datetime import datetime
import ctypes
import subprocess
from ctypes import wintypes

class MainApp:

    class SYSTEMTIME(ctypes.Structure):
        _fields_ = [
            ("wYear", wintypes.WORD),
            ("wMonth", wintypes.WORD),
            ("wDayOfWeek", wintypes.WORD),
            ("wDay", wintypes.WORD),
            ("wHour", wintypes.WORD),
            ("wMinute", wintypes.WORD),
            ("wSecond", wintypes.WORD),
            ("wMilliseconds", wintypes.WORD)
        ]

    def __init__(self, master=None):
        # build ui
        self.mainToplevel = tk.Tk() if master is None else tk.Toplevel(master)
        self.mainToplevel.configure(
            background="#e0e0e0", height=576, width=768)
        self.mainToplevel.geometry("768x576")
        self.mainToplevel.resizable(False, False)
        self.mainToplevel.title("MSC Merge Files in SharePoint (v0.2)")
        self.lbl_username = ttk.Label(self.mainToplevel)
        self.lbl_username.configure(background="#e0e0e0", text='Username :')
        self.lbl_username.place(anchor="nw", x=20, y=20)
        self.lbl_password = ttk.Label(self.mainToplevel)
        self.lbl_password.configure(background="#e0e0e0", text='Password :')
        self.lbl_password.place(anchor="nw", x=20, y=40)
        self.ent_username = ttk.Entry(self.mainToplevel)
        self.ent_username.configure(width=30)
        self.ent_username.place(anchor="nw", x=90, y=20)
        self.ent_password = ttk.Entry(self.mainToplevel)
        self.ent_password.configure(show="•", width=30)
        self.ent_password.place(anchor="nw", x=90, y=40)
        self.ent_password.bind("<Return>", self.on_enter_password)
        self.btn_auth = ttk.Button(self.mainToplevel)
        self.btn_auth.configure(text='Login')
        self.btn_auth.place(anchor="nw", x=300, y=30)
        self.btn_auth.configure(command=self.authenticate)
        self.btn_merge = ttk.Button(self.mainToplevel)
        self.btn_merge.configure(text='Merge')
        self.btn_merge.place(anchor="nw", x=50, y=480)
        self.btn_merge.configure(command=self.download_files)
        self.lst_file = tk.Listbox(self.mainToplevel)
        self.lst_file.configure(selectmode="multiple")
        self.lst_file.place(anchor="nw", height=250, width=650, x=50, y=120)
        self.cmb_year = ttk.Combobox(self.mainToplevel)
        self.cmb_year.configure(state="readonly")
        self.cmb_year.place(anchor="nw", width=80, x=90, y=80)
        self.lbl_year = ttk.Label(self.mainToplevel)
        self.lbl_year.configure(background="#e0e0e0", text='Year')
        self.lbl_year.place(anchor="nw", x=50, y=80)
        self.lbl_grade = ttk.Label(self.mainToplevel)
        self.lbl_grade.configure(background="#e0e0e0", text='Grade')
        self.lbl_grade.place(anchor="nw", x=200, y=80)
        self.cmb_grade = ttk.Combobox(self.mainToplevel)
        self.cmb_grade.configure(state="readonly")
        self.cmb_grade.place(anchor="nw", width=80, x=250, y=80)
        self.lbl_task = ttk.Label(self.mainToplevel)
        self.lbl_task.configure(background="#e0e0e0", text='Task')
        self.lbl_task.place(anchor="nw", x=370, y=80)
        self.cmb_task = ttk.Combobox(self.mainToplevel)
        self.cmb_task.configure(state="readonly")
        self.cmb_task.place(anchor="nw", width=80, x=410, y=80)
        self.btn_list = ttk.Button(self.mainToplevel)
        self.btn_list.configure(text='List Files')
        self.btn_list.place(anchor="nw", x=520, y=75)
        self.btn_list.configure(command=self.list_files)
        self.ent_saveas = ttk.Entry(self.mainToplevel)
        self.ent_saveas.configure(width=70)
        self.ent_saveas.place(anchor="nw", width=540, x=50, y=400)
        self.lbl_output = ttk.Label(self.mainToplevel)
        self.lbl_output.configure(
            background="#e0e0e0",
            text='Output location :')
        self.lbl_output.place(anchor="nw", x=50, y=380)
        self.btn_change = ttk.Button(self.mainToplevel)
        self.btn_change.configure(text='Change location')
        self.btn_change.place(anchor="nw", x=600, y=395)
        self.btn_change.configure(command=self.change_location)
        self.ent_filename = ttk.Entry(self.mainToplevel)
        self.ent_filename.configure(width=70)
        self.ent_filename.place(anchor="nw", width=540, x=50, y=450)
        self.ent_filename.bind("<Return>", self.on_enter_merge)
        self.lbl_filename = ttk.Label(self.mainToplevel)
        self.lbl_filename.configure(
            background="#e0e0e0",
            text='Output filename :')
        self.lbl_filename.place(anchor="nw", x=50, y=430)
        self.lbl_loading_bg = tk.Label(self.mainToplevel)
        self.lbl_loading = tk.Label(self.mainToplevel)
        #self.load_gif()

        # set output folder
        documents_folder = os.path.expanduser("~\\Documents")
        self.ent_saveas.insert(0, documents_folder)
        self.ent_saveas.configure(state="readonly")

        # set username and password
        self.ent_username.insert(0, self.get_email())
        #self.ent_password.insert(0, "")
        self.init_input()

        # Main widget
        self.mainwindow = self.mainToplevel

    def run(self):
        self.mainwindow.mainloop()

    def get_email(self):
        command = "whoami /upn"

        try:
            output = subprocess.check_output(command, shell=True, encoding="utf-8")
            return output.strip()
        except subprocess.CalledProcessError as e:
            print(f"Error executing command: {e}")
            return None

    def on_enter_password(self, event):
        self.authenticate()

    def on_enter_merge(self, event):
        self.download_files()

    def remove_unsupported_characters(self, file_name):
        unsupported_chars = r'[<>:"\'/\\|?*]'
        cleaned_name = re.sub(unsupported_chars, '', file_name)
        return cleaned_name

    def resource_path(self, relative_path):
        try:
            base_path = sys._MEIPASS
        except Exception:
            base_path = os.path.abspath(".")

        return os.path.join(base_path, relative_path)

    def set_output_filename(self):
        year = self.cmb_year.get()
        grade = self.cmb_grade.get()
        task = self.cmb_task.get()
        match grade:
            case 'Year 7':
                str_grade = 'Y7'
            case 'Year 8':
                str_grade = 'Y8'
            case 'Year 9':
                str_grade = 'Y9'
            case 'Year 10':
                str_grade = 'Y10'
            case 'Year 11':
                str_grade = 'Y11'
            case 'Year 12':
                str_grade = 'Y12'
        match task:
            case'Task 1':
                str_task = 'AT1N'
            case 'Task 2':
                str_task = 'AT2N'
            case 'Task 3':
                str_task = 'AT3N'
            case 'Task 4':
                str_task = 'AT4N'

        out_filename = f"{year}_{str_grade}_{str_task}"
        #print(out_filename)
        self.ent_filename.delete(0, tk.END)
        self.ent_filename.insert(0, out_filename)
        return out_filename

    def load_gif(self):
        #disable button
        self.btn_list['state'] = 'disabled'
        self.btn_change['state'] = 'disabled'
        self.btn_merge['state'] = 'disabled'
        self.cmb_year['state'] = 'disabled'
        self.cmb_grade['state'] = 'disabled'
        self.cmb_task['state'] = 'disabled'
        self.ent_filename['state'] = 'disabled'

        # Load the animated GIF
        gif_path = self.resource_path("loading.gif")
        gif_image = Image.open(gif_path)

        # Create a list to store each frame of the animated GIF
        frames = []

        # Split the animated GIF into individual frames
        try:
            while True:
                frames.append(ImageTk.PhotoImage(gif_image))
                gif_image.seek(len(frames))
        except EOFError:
            pass

        # Create a Label and display the first frame of the animated GIF
        #self.lbl_loading_bg = tk.Label(self.mainToplevel)
        self.lbl_loading_bg.configure(
            background="#ffffff",
            justify="left",
            takefocus=False)
        self.lbl_loading_bg.place(
            anchor="nw", height=250, width=650, x=50, y=120)
        #self.lbl_loading = tk.Label(self.mainToplevel)
        self.lbl_loading.configure(
            background="#ffffff",
            cursor="arrow",
            font="TkDefaultFont",
            takefocus=False)
        self.lbl_loading.place(anchor="nw", height=150, x=150, y=170)

        # Function to update the image with the next frame
        def update_image(frame_index):
            frame = frames[frame_index]
            self.lbl_loading.config(image=frame)

            # Schedule the next frame update
            self.mainToplevel.after(50, update_image, (frame_index + 1) % len(frames))

        # Start the animation in a separate thread
        def start_animation():
            #print("running gif")
            update_image(0)

        threading.Thread(target=start_animation, daemon=True).start()

    def unload_gif(self):
        self.lbl_loading_bg.place_forget()
        self.lbl_loading.place_forget()

        #enable button
        self.btn_list['state'] = 'normal'
        self.btn_change['state'] = 'normal'
        self.btn_merge['state'] = 'normal'
        self.cmb_year['state'] = 'readonly'
        self.cmb_grade['state'] = 'readonly'
        self.cmb_task['state'] = 'readonly'
        self.ent_filename['state'] = 'normal'

    def init_input(self):

        self.btn_list['state'] = 'disabled'
        self.btn_change['state'] = 'disabled'
        self.btn_merge['state'] = 'disabled'
        self.cmb_year['state'] = 'disabled'
        self.cmb_grade['state'] = 'disabled'
        self.cmb_task['state'] = 'disabled'
        self.ent_filename['state'] = 'disabled'

    def change_location(self):
        #self.load_gif()
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.ent_saveas.configure(state="normal")
            self.ent_saveas.delete(0, tk.END)
            self.ent_saveas.insert(0, folder_path)
            self.ent_saveas.configure(state="readonly")

    def authenticate(self):
        #self.load_gif()
        if self.ent_password.get() == '':
            messagebox.showerror("Authentication Error", "Invalid username and password Please try again")
            return None

        def auth_thread():
            site_url = "https://schoolsnsw.sharepoint.com/sites/MSCCV19STAFF"
            # Replace with the SharePoint folder relative URL
            #folder_url = "%2Fsites/MSCCV19STAFF/Shared%20Documents%2F1%2E%20Maths%2F2023%2FStage%204%2FYear%207%20Maths%2FAssessment%20Tasks%20and%20Schedules%2FTask%201"
            folder_url = "%2Fsites/MSCCV19STAFF/Shared%20Documents"

            # Get username and password from entry fields
            username = self.ent_username.get()
            password = self.ent_password.get()
            #print(username)
            try:
                self.load_gif()
                # Create an authentication context
                auth_ctx = AuthenticationContext(url=site_url)
                # Authenticate with the username and password
                auth_ctx.acquire_token_for_user(username, password)

                # Create a client context using the authenticated context
                ctx = ClientContext(site_url, auth_ctx)

                # Get the folder by relative URL
                folder = ctx.web.get_folder_by_server_relative_url(folder_url)
                ctx.load(folder)
                ctx.execute_query()

                # Get the files in the folder
                files = folder.files
                ctx.load(files)
                ctx.execute_query()

                self.ent_username['state'] = 'readonly'
                self.ent_password['state'] = 'readonly'
                self.btn_auth['state'] = 'disabled'

                # Define the value list for the Combobox
                year_values = ["2022", "2023", "2024"]
                self.cmb_year['values'] = year_values
                self.cmb_year.set("2023")
                grade_values = ["Year 7", "Year 8", "Year 9", "Year 10", "Year 11", "Year 12"]
                self.cmb_grade['values'] = grade_values
                self.cmb_grade.set("Year 7")
                task_values = ["Task 1", "Task 2", "Task 3", "Task 4"]
                self.cmb_task['values'] = task_values
                self.cmb_task.set("Task 1")
                self.unload_gif()
                self.ent_filename['state'] = 'disabled'

            except Exception as e:
                print(f"Error, {e}")
                self.unload_gif()
                messagebox.showerror("Authentication Error", "Invalid username and password Please try again")

        threading.Thread(target=auth_thread, daemon=True).start()
        #self.unload_gif()

    def list_files(self):
        #self.load_gif()

        def list_thread():
            site_url = "https://schoolsnsw.sharepoint.com/sites/MSCCV19STAFF"
            folder_url = "%2Fsites/MSCCV19STAFF/Shared%20Documents"
            # Get username and password from entry fields
            username = self.ent_username.get()
            password = self.ent_password.get()
            target_folder = []
            target_folder1 = []
            target_folder2 = []
            target_folder3 = []
            file_paths = []
            lookup_folder = "Assessment Tasks and Schedules"

            # Clear the listbox
            self.lst_file.delete(0, tk.END)

            try:
                self.load_gif()
                # Create an authentication context
                auth_ctx = AuthenticationContext(url=site_url)
                # Authenticate with the username and password
                auth_ctx.acquire_token_for_user(username, password)

                # Create a client context using the authenticated context
                ctx = ClientContext(site_url, auth_ctx)

                # Get the folder by relative URL
                folder = ctx.web.get_folder_by_server_relative_url(folder_url)
                ctx.load(folder)
                ctx.execute_query()

                # Get the collection of folders within the specified folder
                sub_folders = folder.folders
                ctx.load(sub_folders)
                ctx.execute_query()

                # Print the folder names
                for sub_folder in sub_folders:
                    if str(sub_folder.properties["Name"]).startswith("1. "):
                        if self.cmb_grade.get() == 'Year 7' or self.cmb_grade.get() == 'Year 8':
                            tmp_stage = 'Stage 4'
                        elif self.cmb_grade.get() == 'Year 9' or self.cmb_grade.get() == 'Year 10':
                            tmp_stage = 'Stage 5'
                        elif self.cmb_grade.get() == 'Year 11' or self.cmb_grade.get() == 'Year 12':
                            tmp_stage = 'Stage 6'
                        #tmp_str = f'{str(sub_folder.properties["Name"])}/{self.cmb_year.get()}/{self.cmb_grade.get()}'
                        tmp_str = f'{str(sub_folder.properties["Name"])}/{self.cmb_year.get()}/{tmp_stage}'
                        #tmp_str = tmp_str.replace(' ','%20')
                        #tmp_str = tmp_str.replace('.', '%2E')
                        target_folder.append(tmp_str)
                print(target_folder)

                for sub_fol in target_folder:
                    try:
                        # Get the folder by relative URL
                        tmp_url = f'{folder_url}/{sub_fol}'
                        folder = ctx.web.get_folder_by_server_relative_url(tmp_url)
                        ctx.load(folder)
                        ctx.execute_query()
                        # Get the collection of folders within the specified folder
                        sub_folders = folder.folders
                        ctx.load(sub_folders)
                        ctx.execute_query()
                        for sub_folder in sub_folders:
                            if str(sub_folder.properties["Name"]).startswith(self.cmb_grade.get()):
                                #tmp_str = f'{sub_fol}/{str(sub_folder.properties["Name"])}'
                                tmp_str = f'{sub_fol}/{str(sub_folder.properties["Name"])}'
                                target_folder1.append(tmp_str)

                    except ClientRequestException as e:
                        if "404" in str(e):
                            # Folder does not exist, skip the code block
                            continue
                        else:
                            # Handle other exceptions
                            print(f"Error occurred while accessing folder: {e}")

                    except Exception as e:
                        # Handle other exceptions
                        print(f"Error occurred while accessing folder: {e}")

                print(target_folder1)

                for sub_fol in target_folder1:
                    try:
                        # Get the folder by relative URL
                        tmp_url = f'{folder_url}/{sub_fol}'
                        folder = ctx.web.get_folder_by_server_relative_url(tmp_url)
                        ctx.load(folder)
                        ctx.execute_query()
                        # Get the collection of folders within the specified folder
                        sub_folders = folder.folders
                        ctx.load(sub_folders)
                        ctx.execute_query()
                        for sub_folder in sub_folders:
                            if str(sub_folder.properties["Name"]) == lookup_folder:
                                tmp_str = f'{sub_fol}/{str(sub_folder.properties["Name"])}'
                                target_folder3.append(tmp_str)

                    except ClientRequestException as e:
                        if "404" in str(e):
                            # Folder does not exist, skip the code block
                            continue
                        else:
                            # Handle other exceptions
                            print(f"Error occurred while accessing folder: {e}")

                    except Exception as e:
                        # Handle other exceptions
                        print(f"Error occurred while accessing folder: {e}")
                print(target_folder3)
                # looking at the 4th layer folder
                for sub_fol in target_folder3:
                    try:
                        # Get the folder by relative URL
                        tmp_url = f'{folder_url}/{sub_fol}/{self.cmb_task.get()}'
                        #print(tmp_url)
                        folder = ctx.web.get_folder_by_server_relative_url(tmp_url)
                        ctx.load(folder)
                        ctx.execute_query()
                        # Get the collection of folders within the specified folder
                        sub_files = folder.files
                        ctx.load(sub_files)
                        ctx.execute_query()
                        # Iterate over the files in the folder
                        for file in sub_files:
                            if self.cmb_task.get() == 'Task 1':
                                if file.properties["Name"].lower().endswith("at1n.docx"):
                                    # Get the file path and append it to the list
                                    file_path = f'{str(file.properties["ServerRelativeUrl"])}'
                                    file_path = file_path.replace('/sites/MSCCV19STAFF/Shared Documents/','')
                                    file_paths.append(file_path)
                            elif self.cmb_task.get() == 'Task 2':
                                if file.properties["Name"].lower().endswith("at2n.docx"):
                                    # Get the file path and append it to the list
                                    file_path = f'{str(file.properties["ServerRelativeUrl"])}'
                                    file_path = file_path.replace('/sites/MSCCV19STAFF/Shared Documents/', '')
                                    file_paths.append(file_path)
                            elif self.cmb_task.get() == 'Task 3':
                                if file.properties["Name"].lower().endswith("at3n.docx"):
                                    # Get the file path and append it to the list
                                    file_path = f'{str(file.properties["ServerRelativeUrl"])}'
                                    file_path = file_path.replace('/sites/MSCCV19STAFF/Shared Documents/', '')
                                    file_paths.append(file_path)
                            elif self.cmb_task.get() == 'Task 4':
                                if file.properties["Name"].lower().endswith("at4n.docx"):
                                    # Get the file path and append it to the list
                                    file_path = f'{str(file.properties["ServerRelativeUrl"])}'
                                    file_path = file_path.replace('/sites/MSCCV19STAFF/Shared Documents/', '')
                                    file_paths.append(file_path)

                    except ClientRequestException as e:
                        if "404" in str(e):
                            # Folder does not exist, skip the code block
                            continue
                        else:
                            # Handle other exceptions
                            print(f"Error occurred while accessing folder: {e}")

                    except Exception as e:
                        # Handle other exceptions
                        print(f"Error occurred while accessing folder: {e}")

                def get_file_name(file_path):
                    return file_path.split('/')[-1].split('.')[0]

                sorted_list = sorted(file_paths, key=get_file_name)
                print(sorted_list)

                for path in sorted_list:
                    self.lst_file.insert(tk.END, path)

                self.unload_gif()
                self.set_output_filename()

            except Exception as e:
                messagebox.showerror("Error", str(e))

        threading.Thread(target=list_thread, daemon=True).start()
        #self.unload_gif()

    def download_files(self):
        #self.load_gif()
        if self.ent_filename.get() == '':
            messagebox.showerror("Filename Error", "Please specify the merged filename")
            return None
        cleaned_name = self.remove_unsupported_characters(self.ent_filename.get())
        out_file = f'{cleaned_name}.docx'

        def merge_thread():
            temp_folder = os.environ.get('TEMP')
            #print(f"User's local temporary folder: {temp_folder}")
            tmp_files = []

            destination_folder = self.ent_saveas.get()
            all_files = self.lst_file.get(0, tk.END)

            # Replace with your SharePoint site URL
            site_url = "https://schoolsnsw.sharepoint.com/sites/MSCCV19STAFF"
            # Replace with the SharePoint folder relative URL
            #folder_url = "%2Fsites/MSCCV19STAFF/Shared%20Documents%2F1%2E%20Maths%2F2023%2FStage%204%2FYear%207%20Maths%2FAssessment%20Tasks%20and%20Schedules%2FTask%201"
            folder_url = "/sites/MSCCV19STAFF/Shared Documents"

            # Get username and password from entry fields
            username = self.ent_username.get()
            password = self.ent_password.get()

            def remove_headers_and_footers(doc):
                for section in doc.sections:
                    for header in section.header.paragraphs:
                        header.text = ""
                    for footer in section.footer.paragraphs:
                        footer.text = ""

            def combine_all_docx(filename_master, files_list):
                number_of_sections = len(files_list)
                master = Document_compose(filename_master)
                composer = Composer(master)
                for i in range(0, number_of_sections):
                    doc_temp = Document_compose(files_list[i])
                    composer.append(doc_temp)
                    if i < number_of_sections - 1:
                        master.add_page_break()
                composer.save(f'{destination_folder}/{out_file}')
                # Open the existing DOCX file
                doc = Document(f'{destination_folder}/{out_file}')
                # Access the core properties of the document
                core_props = doc.core_properties
                # Set the "date last saved" property to the current date and time
                kernel32 = ctypes.WinDLL('kernel32', use_last_error=True)
                GetSystemTime = kernel32.GetSystemTime
                GetSystemTime.argtypes = [ctypes.POINTER(self.SYSTEMTIME)]
                GetSystemTime.restype = None
                system_time = self.SYSTEMTIME()
                GetSystemTime(ctypes.byref(system_time))
                current_time = datetime(
                    system_time.wYear,
                    system_time.wMonth,
                    system_time.wDay,
                    system_time.wHour,
                    system_time.wMinute,
                    system_time.wSecond,
                    system_time.wMilliseconds
                )
                #print(current_time)
                # # Remove headers and footers from the document
                # remove_headers_and_footers(doc)
                # Save the modified document to a new file
                core_props.modified = current_time
                doc.save(f'{destination_folder}/{out_file}')

            try:
                self.load_gif()
                # Create an authentication context
                auth_ctx = AuthenticationContext(url=site_url)
                # Authenticate with the username and password
                auth_ctx.acquire_token_for_user(username, password)

                # Create a client context using the authenticated context
                ctx = ClientContext(site_url, auth_ctx)

                # Download each selected file
                for file_name in all_files:
                    # Get the folder by relative URL
                    directory_path, f_name = os.path.split(file_name)
                    folder = ctx.web.get_folder_by_server_relative_url(folder_url + "/" + directory_path)
                    ctx.load(folder)
                    ctx.execute_query()

                    file = folder.files.get_by_url(f_name)
                    #print(f'{file_name}')
                    ctx.load(file)
                    ctx.execute_query()

                    # Download the file content
                    local_path = f"{temp_folder}/{f_name}"
                    tmp_files.append(local_path)
                    with open(local_path, "wb") as local_file:
                        file.download(local_file).execute_query()

                #messagebox.showinfo("Download Complete", "Selected files have been downloaded.")
                # Split the first item and the rest of the items
                filename_master = tmp_files[0]
                files_list = tmp_files[1:]
                combine_all_docx(filename_master, files_list)

                # Save the combined document to a new file
                self.unload_gif()
                messagebox.showinfo("Merge Completed", f"Merged document have been saved to {destination_folder}\{out_file}")

            except Exception as e:
                messagebox.showerror("Merge Error", str(e))

        threading.Thread(target=merge_thread, daemon=True).start()

if __name__ == "__main__":
    app = MainApp()
    app.run()